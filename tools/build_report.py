from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "report" / "项目报告.md"
OUTPUT = ROOT / "report" / "自动驾驶VLA自进化研究报告.docx"

BLUE = "245B8A"; DARK = "163A5F"; MUTED = "667085"; LIGHT = "EAF1F7"; GRID = "B8C7D6"


def font(run, name="宋体", size=11, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn("w:shd"))
    if shd is None: shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc.get_or_add_tcPr(); m = tc.first_child_found_in("w:tcMar")
    if m is None: m = OxmlElement("w:tcMar"); tc.append(m)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = m.find(qn(f"w:{key}"))
        if node is None: node = OxmlElement(f"w:{key}"); m.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    for tag in ("tblW", "tblInd"):
        old = tblPr.find(qn(f"w:{tag}"))
        if old is not None: tblPr.remove(old)
    tw = OxmlElement("w:tblW"); tw.set(qn("w:w"), "9360"); tw.set(qn("w:type"), "dxa"); tblPr.append(tw)
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tblPr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(w)); grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr(); tcW = tcPr.find(qn("w:tcW"))
            if tcW is None: tcW = OxmlElement("w:tcW"); tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths[i])); tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; margins(cell)


def add_field(paragraph, instruction, fallback=""):
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = fallback
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    for x in (begin, instr, sep, text, end): run.append(x)


def setup_styles(doc):
    normal = doc.styles["Normal"]; normal.font.name = "宋体"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体"); normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(8); normal.paragraph_format.line_spacing = 1.333
    for name, size, before, after, color in (("Heading 1",16,18,10,BLUE),("Heading 2",13,12,6,BLUE),("Heading 3",12,8,4,DARK)):
        st=doc.styles[name]; st.font.name="微软雅黑"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑"); st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    for name in ("List Bullet", "List Number"):
        st=doc.styles[name]; st.font.name="宋体"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体"); st.font.size=Pt(10.5); st.paragraph_format.left_indent=Inches(.375); st.paragraph_format.first_line_indent=Inches(-.194); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.208


def add_inline(p, text, size=10.5, bold=False, color=None):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part: continue
        code = part.startswith("`") and part.endswith("`")
        r = p.add_run(part[1:-1] if code else part)
        font(r, "等线" if code else "宋体", 9.5 if code else size, bold=bold or code, color=color)


def add_table(doc, rows):
    cols = len(rows[0]); table = doc.add_table(rows=len(rows), cols=cols)
    widths = [9360 // cols] * cols; widths[-1] += 9360 - sum(widths)
    if cols == 3: widths = [2250, 5550, 1560]
    if cols == 4: widths = [2100, 3150, 2550, 1560]
    if cols == 8: widths = [1500, 1050, 1050, 1050, 1100, 1100, 1200, 1310]
    set_table_geometry(table, widths)
    for i,row in enumerate(rows):
        trPr=table.rows[i]._tr.get_or_add_trPr(); cant=OxmlElement("w:cantSplit"); trPr.append(cant)
        for j,text in enumerate(row):
            text = text.replace("`", "")
            c=table.cell(i,j); c.text=""; p=c.paragraphs[0]
            p.paragraph_format.space_after=Pt(0 if cols == 3 else 2)
            p.paragraph_format.line_spacing=1.0 if cols == 3 else 1.05
            if i == 0: p.paragraph_format.keep_with_next = True
            align_center = i == 0 or (cols >= 5 and j > 0) or j == cols-1
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
            body_size = 8.2 if cols >= 5 else (8.5 if cols == 3 else 9.0)
            r=p.add_run(text); font(r,"微软雅黑" if i==0 else "宋体",body_size,bold=i==0,color=DARK if i==0 else None)
            if i==0: shade(c,LIGHT)
    # repeat header
    trPr=table.rows[0]._tr.get_or_add_trPr(); h=OxmlElement("w:tblHeader"); h.set(qn("w:val"),"true"); trPr.append(h)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)


def cover(doc, title_lines):
    sec=doc.sections[0]
    sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    for _ in range(4): doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("课程实践 · 研究报告"); font(r,"微软雅黑",11,True,BLUE); p.paragraph_format.space_after=Pt(18)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(title_lines[0]); font(r,"微软雅黑",25,True,DARK); p.paragraph_format.space_after=Pt(12)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Driving Reflection × Multi-dimensional Critic × Self-Evolution"); font(r,"Calibri",13,False,MUTED); p.paragraph_format.space_after=Pt(54)
    meta=[("课程选题","大区治理选题第 4 题"),("冻结基座","OpenDriveVLA-0.5B（本地已核验）"),("实验版本","轻量自进化 Demo v0.3"),("完成日期","2026 年 8 月")]
    t=doc.add_table(rows=len(meta),cols=2); set_table_geometry(t,[2400,6960])
    for i,(a,b) in enumerate(meta):
        for j,val in enumerate((a,b)):
            c=t.cell(i,j); c.text=""; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(val); font(r,"微软雅黑" if j==0 else "宋体",10.5,j==0,DARK if j==0 else None)
            if j==0: shade(c,"F3F6F9")
    doc.add_page_break()
    p=doc.add_paragraph(); p.style="Heading 1"; p.add_run("目录")
    toc=doc.add_paragraph(); add_field(toc,'TOC \\o "1-3" \\h \\z \\u',"右键更新目录")


def parse_body(doc, text):
    lines=text.splitlines(); i=0
    while i<len(lines):
        line=lines[i].strip()
        if not line: i+=1; continue
        if line.startswith("|"):
            rows=[]
            while i<len(lines) and lines[i].strip().startswith("|"):
                vals=[x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?",x) for x in vals): rows.append(vals)
                i+=1
            add_table(doc,rows); continue
        m=re.match(r"^(#{2,4})\s+(.+)$",line)
        if m:
            level=len(m.group(1))-1
            p=doc.add_paragraph(style=f"Heading {min(level,3)}"); add_inline(p,m.group(2),size={1:16,2:13,3:12}.get(level,11),bold=True,color=BLUE if level<3 else DARK); i+=1; continue
        if re.match(r"^\d+\.\s+",line):
            p=doc.add_paragraph(style="List Number"); add_inline(p,re.sub(r"^\d+\.\s+","",line)); i+=1; continue
        if line.startswith("- "):
            p=doc.add_paragraph(style="List Bullet"); add_inline(p,line[2:]); i+=1; continue
        if re.match(r"^\[\d+\]\s+", line):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.28); p.paragraph_format.first_line_indent=Inches(-.28); p.paragraph_format.space_after=Pt(5); add_inline(p,line,size=9.5); i+=1; continue
        para=[line]; i+=1
        while i<len(lines) and lines[i].strip() and not re.match(r"^(#{2,4})\s+|^\d+\.\s+|^-\s+|^\|",lines[i].strip()):
            para.append(lines[i].strip()); i+=1
        joined = " ".join(para)
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT if "http" in joined else WD_ALIGN_PARAGRAPH.JUSTIFY; add_inline(p,joined)


def main():
    text=SOURCE.read_text(encoding="utf-8")
    title=text.splitlines()[0].lstrip("# ")
    marker="## 摘要"; body=marker+text.split(marker,1)[1]
    doc=Document(); setup_styles(doc); cover(doc,[title]); parse_body(doc,body)
    sec=doc.sections[0]
    sec.different_first_page_header_footer=True
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; r=header.add_run("自动驾驶 VLA 自进化研究"); font(r,"微软雅黑",9,False,MUTED)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run("— "); font(r,"Calibri",9,color=MUTED); add_field(footer,"PAGE","1"); r=footer.add_run(" —"); font(r,"Calibri",9,color=MUTED)
    sec.first_page_header.paragraphs[0].text=""; sec.first_page_footer.paragraphs[0].text=""
    settings=doc.settings._element; update=OxmlElement("w:updateFields"); update.set(qn("w:val"),"true"); settings.append(update)
    props=doc.core_properties; props.title=title; props.subject="课程实践选题4项目报告"; props.author="课程项目组"; props.keywords="VLA, Driving Reflection, Critic, Self-Evolution"
    OUTPUT.parent.mkdir(parents=True,exist_ok=True); doc.save(OUTPUT); print(OUTPUT)


if __name__=="__main__": main()
